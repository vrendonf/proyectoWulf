import sounddevice as sd
import numpy as np
import queue
import whisper
import threading
import time
import soundfile as sf
import os
import torch
import customtkinter as ctk
from deep_translator import GoogleTranslator  # Usamos deep_translator nuevamente
from gtts import gTTS
from datetime import datetime
from docx import Document

# 🚀 Configuración
SAMPLE_RATE = 16000
CHANNELS = 1
FILENAME = "temp_audio.wav"

# 🔹 Variables globales
audio_queue = queue.Queue()
recording = False
mic_stream = None
doc = Document()
last_translated_text = ""

# 🚀 Cargar Whisper en GPU si está disponible
device = "cuda" if torch.cuda.is_available() else "cpu"
model = whisper.load_model("small").to(device)  # Se puede cambiar a "tiny", "medium"

# 🔹 Obtener dispositivos de audio
def get_audio_devices():
    devices = sd.query_devices()
    return [(i, dev["name"]) for i, dev in enumerate(devices) if dev["max_input_channels"] > 0]

# 🎤 Callback de grabación en streaming
def mic_callback(indata, *_):
    if recording:
        audio_queue.put(indata.copy())

# 📄 Guardar texto en documento de Word
def save_to_word(original_text, translated_text):
    global doc
    current_time = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
    doc.add_paragraph(f"Fecha y hora: {current_time}")
    doc.add_heading('Texto Original', level=2)
    doc.add_paragraph(original_text)
    doc.add_heading('Traducción', level=2)
    doc.add_paragraph(translated_text)
    doc.add_paragraph("-----------------------------")
    doc.save("transcripcion_traduccion.docx")

# 🚀 Procesar audio en tiempo real con más precisión
def process_audio():
    global last_translated_text
    while True:
        if not audio_queue.empty():
            audio_data = []
            while not audio_queue.empty():
                audio_data.append(audio_queue.get())
            
            if audio_data:
                # Guardar audio con soundfile
                audio_array = np.concatenate(audio_data, axis=0).astype(np.float32)
                sf.write(FILENAME, audio_array, SAMPLE_RATE)

                # 🎙️ Transcripción con Whisper
                result = model.transcribe(FILENAME, fp16=False, temperature=0.1)
                transcribed_text = result["text"].strip()

                # 🌍 Detección de idioma con Whisper
                detected_lang = result["language"]

                # 🌍 Traducción automática con `deep_translator`
                target_lang = "es" if detected_lang != "es" else "en"
                translator = GoogleTranslator(source=detected_lang, target=target_lang)
                translated_text = translator.translate(transcribed_text) if transcribed_text else "⚠️ No se detectó texto."

                # Evitar traducciones innecesarias
                if transcribed_text.lower() != translated_text.lower():
                    # 🖥️ Mostrar en la interfaz
                    text_box.insert("end", f"🎧 Original ({detected_lang}): {transcribed_text}\n")
                    text_box.insert("end", f"🌍 Traducción: {translated_text}\n")
                    text_box.insert("end", "-----------------------------\n")
                    text_box.yview("end")

                    # 🗣️ Guardar última traducción para reproducción de voz
                    last_translated_text = translated_text

                    # 📄 Guardar en documento Word
                    save_to_word(transcribed_text, translated_text)
                else:
                    print("🔹 [INFO] Traducción similar al original, descartada.")

        time.sleep(0.1)  # 🔹 Captura más datos antes de procesar

# 🎤 Iniciar grabación de audio
def start_recording():
    global recording, mic_stream
    stop_recording()  # 🔹 Asegurar que no haya otra grabación activa
    
    recording = True
    selected_device_name = device_combobox.get()  # Obtener el nombre del dispositivo
    selected_device = next((i for i, name in audio_devices if name == selected_device_name), None)

    if selected_device is not None:
        device_index = audio_devices[selected_device][0]
        mic_stream = sd.InputStream(callback=mic_callback, samplerate=SAMPLE_RATE, channels=CHANNELS, device=device_index)
        mic_stream.start()
        print("🎤 Grabando...")
    else:
        print("⚠️ No se seleccionó un dispositivo de entrada.")
        recording = False

# ⏹️ Detener grabación
def stop_recording():
    global recording, mic_stream
    if mic_stream:
        mic_stream.stop()
        mic_stream.close()
    recording = False

# 🔊 Reproducir la traducción con voz
def speak_text():
    global last_translated_text
    if last_translated_text:
        tts = gTTS(last_translated_text, lang="es")
        tts.save("translated_audio.mp3")
        os.system("afplay translated_audio.mp3")
    else:
        print("⚠️ No hay traducción para reproducir.")

# 📄 Guardar texto de la interfaz en Word
def save_current_text_to_word():
    text = text_box.get("1.0", "end")  # Obtener todo el texto de la caja de texto
    doc.add_paragraph(text)
    doc.save("transcripcion_traduccion.docx")
    print("📄 Texto guardado en 'transcripcion_traduccion.docx'")

# ---- 🎨 Interfaz gráfica optimizada con customtkinter ----
ctk.set_appearance_mode("dark")
ctk.set_default_color_theme("blue")

root = ctk.CTk()
root.title("🎙️ Wulf Translate")
root.geometry("600x500")

# 📝 Caja de texto para mostrar transcripciones
text_box = ctk.CTkTextbox(root, wrap="word", width=550, height=250)
text_box.pack(pady=10)

# 🎤 Obtener dispositivos de audio
audio_devices = get_audio_devices()
device_names = [name for _, name in audio_devices]

device_label = ctk.CTkLabel(root, text="Selecciona el dispositivo de audio:")
device_label.pack(pady=5)

device_combobox = ctk.CTkComboBox(root, values=device_names)
device_combobox.pack(pady=5)
device_combobox.set(device_names[0] if device_names else "No disponible")

# 🎙️ Botón para iniciar grabación
start_button = ctk.CTkButton(root, text="🎙️ Iniciar Grabación", command=start_recording)
start_button.pack(pady=5)

# ⏹️ Botón para detener grabación
stop_button = ctk.CTkButton(root, text="⏹️ Detener Grabación", command=stop_recording)
stop_button.pack(pady=5)

# 🔊 Botón para reproducir la traducción en voz
speak_button = ctk.CTkButton(root, text="🔊 Reproducir Voz", command=speak_text)
speak_button.pack(pady=5)

# 📄 Botón para guardar en Word
save_word_button = ctk.CTkButton(root, text="💾 Guardar en Word", command=save_current_text_to_word)
save_word_button.pack(pady=5)

# 🧵 Hilo para procesar audio en segundo plano
thread = threading.Thread(target=process_audio, daemon=True)
thread.start()

# 🎨 Ejecutar interfaz
root.mainloop()
